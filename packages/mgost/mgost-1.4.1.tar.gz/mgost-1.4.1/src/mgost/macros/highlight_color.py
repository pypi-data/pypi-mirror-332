from docx.enum.text import WD_COLOR_INDEX

from . import logger
from ._flags import MacrosFlags
from ._mixins import DuringDocxCreation


class Macros(DuringDocxCreation):
    """Changes background color of run. Use WD_COLOR_INDEX names"""
    __slots__ = ()

    def __init__(self, macros) -> None:
        super().__init__(macros)
        self.macros = macros

    def process_during_docx_creation(self, p, context):
        if len(self.macros.args) != 1:
            logger.info(
                f'Macros "{self.get_name()}" requires at least one argument'
            )
            return [p.add_run("<Arguments error>")]
        name = self.macros.args[0]
        names = {i.name for i in WD_COLOR_INDEX}
        if name not in names:
            logger.info(
                f'Macros "{self.get_name()}" second argument is'
                f' value from WD_COLOR_INDEX. But {name[:30]}'
                f' does not element of {names}'
            )
            return [p.add_run("<Arguments error>")]
        run = p.add_run(self.macros.value)
        run.font.highlight_color = getattr(WD_COLOR_INDEX, self.macros.args[0])

        return [run]

    @staticmethod
    def flags():
        return MacrosFlags.NONE
