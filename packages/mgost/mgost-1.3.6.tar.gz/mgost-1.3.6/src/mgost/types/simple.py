from pathlib import Path
from typing import Sequence

from docx.enum.text import WD_ALIGN_PARAGRAPH

from .functions import add_run, init_paragraph
from .low_functions import add_bookmark, remove_paragraph
from .media import Image
from .mixins import AddableToDocument, AddableToParagraph, ListElement
from .run import Run
from mgost.context import Context, ContextVariable

__all__ = ('Root', 'Paragraph', 'Heading')


class Root(ListElement[AddableToDocument], AddableToDocument):
    __slots__ = ('file_path',)
    file_path: Path | None

    def __init__(self, elements: Sequence[AddableToDocument]) -> None:
        super().__init__()
        assert isinstance(elements, list)
        assert all((isinstance(i, AddableToDocument) for i in elements))
        self.elements = elements
        self.file_path = None

    def __repr__(self) -> str:
        path = ''
        if self.file_path is not None:
            path = f" from {self.file_path}"
        return f"<Root of {len(self.elements)} elements{path}>"

    def add_to_document(self, d, context) -> None:
        for el in self.elements:
            assert isinstance(el, AddableToDocument)
            el.add_to_document(d, context)

    def as_dict(self):
        return {
            f'{type(self).__qualname__}': [
                i.as_dict() for i in self.elements
            ]
        }

    def walk(self):
        yield (self,)
        for element in self.elements:
            for sub_elements in element.walk():
                yield (self, *sub_elements)


class Paragraph(ListElement[AddableToParagraph], AddableToDocument):
    __slots__ = ('do_not_add_tab', 'no_new_line',)

    def __init__(self, runs: list[AddableToParagraph]):
        super().__init__()
        assert isinstance(runs, list)
        self.do_not_add_tab: bool = False
        self.no_new_line: bool | None = None
        self.add(*runs)

    def __repr__(self) -> str:
        return f"<Paragraph of {len(self.elements)} runs>"

    def add_to_document(self, d, context) -> None:
        if not self:
            return
        with (
            ContextVariable(
                context, 'no_new_line', self.no_new_line,
                apply=self.no_new_line is not None
            )
        ):
            if not self.elements:
                return
            p = init_paragraph(d.add_paragraph(), context)
            if 'do_not_add_tab' not in context and not self.do_not_add_tab:
                add_run('\t', p, context)
            having_image = False
            sub_runs = []
            for run in self.elements:
                if isinstance(run, Image):
                    having_image = True
                run = run.add_to_paragraph(p, context)
                sub_runs.extend(run)
            if having_image:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not sub_runs:
                remove_paragraph(d, -1)

    def add_to_paragraph(self, p, context):
        _runs = []
        for run in self.elements:
            _runs.append(run.add_to_paragraph(p, context))
        return _runs

    def as_dict(self):
        output = dict()
        if self.do_not_add_tab:
            output['do_not_add_tab'] = str(self.do_not_add_tab)
        if self.no_new_line is not None:
            output['no_new_line'] = str(self.no_new_line)
        output[f'{type(self).__qualname__}'] = [
            i.as_dict() for i in self.elements
        ]
        return output

    def walk(self):
        yield (self,)
        for element in self.elements:
            for sub_elements in element.walk():
                yield (self, *sub_elements)


class Heading(ListElement[AddableToParagraph], AddableToDocument):
    __slots__ = ('level', '_centered', '_bookmark_name',)

    def __init__(self, elements: list[AddableToParagraph], level: int):
        super().__init__()
        assert isinstance(elements, list)
        assert isinstance(level, int) and (1 <= level <= 3), level
        self.add(*elements)
        self.level = level
        self._bookmark_name = None
        self._centered = False

    def to_headings_counter(self) -> list[int] | None:
        assert isinstance(self.elements[0], Run)
        start = self.elements[0].start
        if start is None:
            return
        if not start[0].isdigit():
            return
        start = start[:start.find(' ')]
        return [int(i) for i in start.split('.')]

    def update_headings_counter(self, context: Context) -> None:
        counters = self.to_headings_counter()
        if counters:
            context.counters.headings = counters

    def add_to_document(self, d, context) -> None:
        if not self:
            return
        assert isinstance(self.elements[0], Run)
        p = d.add_heading(level=self.level)
        self.update_headings_counter(context)
        style_name = f"Heading {self.level}"
        if self._centered:
            style_name = style_name + 'C'
        p.style = style_name
        if self._bookmark_name is None:
            self._bookmark_name = f"_{self.elements[0].start}"
        add_bookmark(p, context, w_id_str=self._bookmark_name)
        context.counters.bookmarks[self._bookmark_name] = self
        for run in self.elements:
            run.add_to_paragraph(p, context)

    def as_dict(self):
        return {
            f'{type(self).__qualname__} {self.level}': [
                i.as_dict() for i in self.elements
            ]
        }
