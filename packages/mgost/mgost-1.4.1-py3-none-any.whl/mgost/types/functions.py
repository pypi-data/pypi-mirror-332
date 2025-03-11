from typing import Protocol

from docx.text.paragraph import Paragraph as _Paragraph
from docx.text.run import Run as _Run
from lxml import etree

from .low_functions import add_hyperlink, add_linked_run
from mgost.context import Context

__all__ = ('init_paragraph', 'add_run',)


MATHML_STRING = '''
<math xmlns="http://www.w3.org/1998/Math/MathML">{}</math>
'''.strip()


class CodeProcessor(Protocol):
    def __call__(
        self,
        p: _Paragraph,
        code: str,
        context: Context,
        *args: str
    ) -> _Run: ...


def init_paragraph(p: _Paragraph, context: Context) -> _Paragraph:
    assert len(p.runs) == 0
    if (marker := context.get('marker', False)):
        p.add_run(marker).add_tab()
    if (style := context.get('paragraph_style', False)):
        p.style = style
    return p


def add_run(text: str | None, p: _Paragraph, context: Context) -> _Run | None:
    if text is None:
        return p.add_run()
    assert isinstance(text, str)
    assert text
    if context.get('no_new_line', True):
        text = text.strip('\n')
    href: str = context.get('href', '')
    assert isinstance(href, str)
    if not href:
        run = p.add_run(text)
    elif href.startswith('http'):
        assert isinstance(href, str)
        run = add_hyperlink(p, text, href)
    else:
        run = add_linked_run(p, text, href, context)
    if context.get('strike', False): run.font.strike = True
    if context.get('bold', False): run.bold = True
    if context.get('italic', False): run.italic = True
    if context.get('strip', False): run.text = run.text.strip()
    return run


def add_formula(p: _Paragraph, text: str, /, context: Context) -> None:
    from latex2mathml.converter import convert as latex2mathml

    math_mls = [latex2mathml(i) for i in text.split('=')]

    math_omml = '<mo>=</mo>'.join(math_mls)

    string = MATHML_STRING.format(math_omml)
    tree = etree.fromstring(string)  # type: ignore
    new_dom = context.mml2omml_xslt()(tree)
    p._p.append(new_dom.getroot())
