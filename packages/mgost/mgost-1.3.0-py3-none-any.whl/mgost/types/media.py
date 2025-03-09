from abc import abstractmethod
from io import BytesIO, StringIO
from pathlib import Path
from typing import TYPE_CHECKING

import PIL.Image
from docx.document import Document as _Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from .functions import add_formula, add_run, init_paragraph
from .mixins import AddableToDocument, AddableToParagraph, HasText, ListElement
from mgost.context import Context, ContextVariable
from mgost.exceptions import VariablesConflict

if TYPE_CHECKING:
    from .simple import Paragraph
    from mgost.context import Counters


class BaseMedia(ListElement[HasText], AddableToDocument):
    __slots__ = ()

    def initialize_variable(
        self,
        title: str,
        context: Context,
        increase_counter: bool = True
    ) -> tuple[int, int]:
        assert isinstance(title, str)
        assert isinstance(context, Context)
        assert isinstance(increase_counter, bool)

        if title in context:
            raise VariablesConflict(f"Переменная {title} объявляется повторно")
        context[title] = self

        counter_cl = self.counter(context.counters)
        heading_num = context.counters.headings[0]
        if heading_num not in counter_cl:
            counter_cl[heading_num] = 1
        counter = counter_cl[heading_num]
        if increase_counter:
            counter_cl[heading_num] += 1

        return heading_num, counter

    @classmethod
    def mention(
        cls,
        headings: list[int],
        counter: dict[int, int]
    ) -> str:
        return (
            f"{cls.short_name()} "
            f"{headings[0]}.{counter.get(headings[0], 0)+1}"
        )

    @staticmethod
    @abstractmethod
    def counter(counters: 'Counters') -> dict[int, int]:
        ...

    @staticmethod
    @abstractmethod
    def short_name() -> str:
        ...


class Image(BaseMedia):
    __slots__ = ('src', 'title')

    def __init__(
        self,
        src: str,
        alt: list[HasText],
        title: str | None = None
    ) -> None:
        super().__init__()
        assert isinstance(src, str)
        assert isinstance(alt, list)
        assert title is None or isinstance(title, str)
        assert all(isinstance(i, HasText) for i in alt)
        assert all(isinstance(i, AddableToParagraph) for i in alt)
        self.src = src
        self.elements = alt
        self.title = title

    def resize_image(
        self,
        path: Path,
        context: Context
    ) -> BytesIO:
        img = PIL.Image.open(path)
        image_sizes = img.size
        sec = context.d.sections[-1]
        page_sizes = (
            sec.page_width.inches - (  # type:ignore
                sec.left_margin.inches + sec.right_margin.inches  # type:ignore
            ),
            sec.page_height.inches - (  # type:ignore
                sec.top_margin.inches + sec.bottom_margin.inches  # type:ignore
            )
        )

        page_sizes = tuple(i*70 for i in page_sizes)
        ratio = tuple(image_sizes[i] / page_sizes[i] for i in range(2))
        scaling_factor = max(*ratio)
        new_size = tuple(int(i / scaling_factor) for i in image_sizes)
        assert len(new_size) == 2
        img = img.resize(new_size)

        img_byte_arr = BytesIO()
        img.save(img_byte_arr, format='png')
        return img_byte_arr

    def add_to_document(self, d: _Document, context: Context) -> None:
        with (
            ContextVariable(context, 'paragraph_style', 'Image')
        ):
            path = context.source.parent / self.src
            if not path.exists():
                print(
                    f"File {path} does not exist. "
                    "Can't place this image in document"
                )
                return

            image = self.resize_image(path, context)
            title = ''.join([str(i.get_text()) for i in self.elements])
            counter_h, counter_curr = self.initialize_variable(
                title, context, increase_counter=True
            )

            subtitle = StringIO()
            subtitle.write('\nРис ')
            subtitle.write(str(counter_h))
            subtitle.write('.')
            subtitle.write(str(counter_curr))
            subtitle.write(' - ')

            init_paragraph(
                d.add_paragraph(),
                context
            ).add_run().add_picture(image)
            p = init_paragraph(d.add_paragraph(), context)
            add_run(subtitle.getvalue(), p, context)
            for alt_run in self.elements:
                assert isinstance(alt_run, AddableToParagraph)
                alt_run.add_to_paragraph(p, context)

    @staticmethod
    def counter(counters) -> dict[int, int]:
        return counters.image

    @staticmethod
    def short_name() -> str:
        return 'рис.'

    def __repr__(self) -> str:
        return (
            f"<Image src={self.src}, "
            f"alt={self.elements}, "
            f"variable={self.title}>"
        )

    def as_dict(self):
        return {
            'src': self.src,
            'alt': [i.as_dict() for i in self.elements],
            'title': self.title
        }


class Table(BaseMedia):
    __slots__ = ('table',)

    def __init__(self, table: list[list['Paragraph']]) -> None:
        super().__init__()
        assert isinstance(table, list)
        assert all((isinstance(i, list) for i in table))
        self.table = table

    def add_to_document(self, d: _Document, context: Context) -> None:
        title = context.table_name
        context.table_name = None
        if title is None:
            raise RuntimeError(f"There's no name for table {self.as_dict()}")

        counter_h, counter_curr = self.initialize_variable(
            title, context, increase_counter=True
        )

        subtitle = StringIO()
        subtitle.write('\nТаблица ')
        subtitle.write(str(counter_h))
        subtitle.write('.')
        subtitle.write(str(counter_curr))
        subtitle.write(' — ')
        subtitle.write(title)

        with ContextVariable(context, 'paragraph_style', 'Table Name'):
            add_run(
                subtitle.getvalue(),
                init_paragraph(d.add_paragraph(), context),
                context
            )

        t = d.add_table(
            len(self.table),
            len(self.table[0]),
            style='Table Grid'
        )
        with ContextVariable(context, 'paragraph_style', 'Table Text'):
            for column_idx, column in enumerate(self.table[0]):
                cell = t.cell(0, column_idx)
                p = cell.paragraphs[0]
                init_paragraph(p, context)
                column.add_to_paragraph(p, context)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for row_idx, row in enumerate(self.table[1:], start=1):
                for column_idx, column in enumerate(row):
                    cell = t.cell(row_idx, column_idx)
                    p = cell.paragraphs[0]
                    init_paragraph(p, context)
                    column.add_to_paragraph(p, context)
        t.autofit = True
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        # table_autofit(t)
        d.add_paragraph()

    @staticmethod
    def counter(counters) -> dict[int, int]:
        return counters.table

    @staticmethod
    def short_name() -> str:
        return 'таблица'

    def as_dict(self):  # type: ignore
        return {
            'rows': [[
                column.as_dict() for column in row
            ] for row in self.table],
        }


class Formula(BaseMedia, AddableToParagraph):
    __slots__ = ('text', 'title')
    text: str
    title: str

    def __init__(self, title: str, text: str) -> None:
        super().__init__()
        assert isinstance(title, str)
        assert isinstance(text, str)
        self.title = title
        self.text = text

    def add_to_document(self, d: _Document, context: Context) -> None:
        table = d.add_table(1, 2)
        formula_cell, numbering_cell = table.rows[0].cells
        formula_p = formula_cell.add_paragraph()
        numbering_p = numbering_cell.add_paragraph()

        page_width = d.sections[-1].page_width
        assert page_width is not None
        formula_cell.width = page_width
        numbering_cell.width = Inches(0.2)

        increase_counter = True
        if self.title.startswith('*'):
            self.title = self.title[1:]
            increase_counter = False
        counter_h, counter_curr = self.initialize_variable(
            self.title, context, increase_counter=increase_counter
        )

        add_formula(formula_p, self.text, context)
        numbering_p.add_run(f"({counter_h}.{counter_curr})")
        d.add_paragraph()

    def add_to_paragraph(self, p, context):
        add_formula(p, self.text, context)
        return []

    @staticmethod
    def counter(counters) -> dict[int, int]:
        return counters.formula

    @staticmethod
    def short_name() -> str:
        return 'формула'

    def as_dict(self):
        return dict()
