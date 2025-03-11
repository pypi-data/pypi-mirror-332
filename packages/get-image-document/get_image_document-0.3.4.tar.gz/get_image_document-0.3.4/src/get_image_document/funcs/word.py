from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


def set_a4_size(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    return section


def add_page_numbers(section):
    """Adds page number in the format 'หน้า X จากทั้งหมด Y หน้า' to the header with custom font."""
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add page number text
    run = paragraph.add_run("หน้า ")
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(12)  # Set font size to 12
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")  # Set Thai font for compatibility

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(ns.qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    instrText1 = OxmlElement("w:instrText")
    instrText1.set(ns.qn("xml:space"), "preserve")
    instrText1.text = "PAGE"
    run._r.append(instrText1)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(ns.qn("w:fldCharType"), "end")
    run._r.append(fldChar2)

    run.add_text(" จาก ")

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(ns.qn("w:fldCharType"), "begin")
    run._r.append(fldChar3)

    instrText2 = OxmlElement("w:instrText")
    instrText2.set(ns.qn("xml:space"), "preserve")
    instrText2.text = "NUMPAGES"
    run._r.append(instrText2)

    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(ns.qn("w:fldCharType"), "end")
    run._r.append(fldChar4)

    run.add_text(" หน้า")

    # Set font for the total page part as well
    run = paragraph.add_run()
    run.font.name = "TH Sarabun New"
    run.font.size = Pt(12)  # Set font size to 12
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")  # Set Thai font for compatibility
