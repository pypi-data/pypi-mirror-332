import datetime
import os
import argparse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tqdm import tqdm
from docx.oxml.ns import qn

from get_image_document.funcs.image_getter import get_images_from_folder
from get_image_document.funcs.image_manage import add_image_keeping_ratio
from get_image_document.funcs.word import add_page_numbers, set_a4_size

from PIL import Image


def process_and_add_image(run, img_path, cell_width_inch, cell_height_inch, dpi):
    """Rotate image if landscape, convert to RGB if necessary, and add it to the DOCX cell while maintaining aspect ratio."""
    try:
        with Image.open(img_path) as img:
            # Rotate if the image is in landscape mode
            if img.width > img.height:
                img = img.rotate(90, expand=True)

            # Convert to RGB if the image has an alpha channel (RGBA mode)
            if img.mode == "RGBA":
                img = img.convert("RGB")

            # Save as a temporary JPEG file
            temp_path = img_path + "_rotated.jpg"
            img.save(temp_path, "JPEG")

            # Add image to DOCX while keeping aspect ratio
            add_image_keeping_ratio(run, temp_path, cell_width_inch, cell_height_inch, dpi)

            # Remove temporary file after adding the image
            os.remove(temp_path)

    except Exception as e:
        print(f"Error processing image {img_path}: {e}")


def create_docx_from_images(image_paths, cols_per_row, rows_per_page, height_scale, dpi):
    """Generates a DOCX file with all images and their filenames in a single table, keeping aspect ratio."""
    try:
        if not image_paths:
            raise ValueError("No images found in the specified directory.")

        doc = Document()

        # **Add Title on the First Page**
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(
            f"รายการภาพจำนวน {format(len(image_paths), ',')} ภาพ")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = "TH Sarabun New"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_before = Pt(0)
        title_paragraph.paragraph_format.space_after = Pt(0)

        section = doc.sections[0]
        section = set_a4_size(section)

        add_page_numbers(section)

        # Start with one row to prevent index errors
        table = doc.add_table(rows=1, cols=cols_per_row)

        error_list = []
        row_idx = 0  # Ensure row index starts at 0

        for i in tqdm(range(len(image_paths)), desc="Processing Images"):
            img_path = image_paths[i]
            col_idx = i % cols_per_row

            # Add a new row when all columns are filled
            if col_idx == 0 and i != 0 and row_idx != rows_per_page - 1:
                row = table.add_row()
                row_idx += 1

            # Add page break when reaching the max rows per page
            if i > 0 and i % (cols_per_row * rows_per_page) == 0:
                doc.add_page_break()
                # New table for new page
                table = doc.add_table(rows=1, cols=cols_per_row)
                row_idx = 0  # Reset row index for the new table

            try:
                cell = table.cell(row_idx, col_idx)
                cell_paragraph = cell.paragraphs[0]
                cell_height_inch = 9 / rows_per_page
                cell_width_inch = 6 / cols_per_row

                # Set paragraph spacing to 0 for image container
                cell_paragraph.paragraph_format.space_before = Pt(0)
                cell_paragraph.paragraph_format.space_after = Pt(0)
                cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Horizontal center alignment

                # Add image while maintaining aspect ratio
                run = cell_paragraph.add_run()
                process_and_add_image(
                    run, img_path, cell_width_inch, cell_height_inch * height_scale, dpi)

                # Add filename with dynamic shrinking
                file_name = os.path.basename(img_path)

                filename_paragraph = cell.add_paragraph()
                filename_paragraph.paragraph_format.space_before = Pt(0)
                filename_paragraph.paragraph_format.space_after = Pt(0)
                filename_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-aligned filename

                filename_run = filename_paragraph.add_run(file_name)
                filename_run.font.name = "TH Sarabun New"
                filename_run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")

                # Shrink font size dynamically
                max_font_size = 16
                min_font_size = 7
                char_per_inch = 10  # Estimate character capacity per inch
                max_chars_fit = cell_width_inch * char_per_inch

                if len(file_name) > max_chars_fit:
                    shrink_factor = len(file_name) / max_chars_fit
                    final_font_size = max(
                        min_font_size, max_font_size / shrink_factor)
                else:
                    final_font_size = max_font_size

                filename_run.font.size = Pt(final_font_size)

            except Exception as e:
                error_msg = f"File: {img_path}: {e}"
                error_list.append(error_msg)

        output_doc = f"combined_images_{datetime.datetime.now().strftime('%Y-%m-%d')}.docx"
        doc.save(output_doc)

        if error_list:
            print("Some images could not be processed:")
            for err in error_list:
                print(f" - {err}")

        print(f"DOCX file created successfully: {output_doc}")

    except ValueError as ve:
        print(f"ERROR: {ve}")
    except FileNotFoundError as fnfe:
        print(f"File not found: {fnfe}")
    except Exception as ex:
        print(f"Unexpected error: {ex}")


def main():
    """Main function to find images and create a DOCX file."""
    parser = argparse.ArgumentParser(description="Generate DOCX from images.")
    parser.add_argument("-c", "--cols_per_row", type=int,default=8, help="Number of columns per row in the DOCX.")
    parser.add_argument("-r", "--rows_per_page", type=int,default=8, help="Number of rows per page in the DOCX.")
    parser.add_argument("-d", "--dpi", type=int, default=150,help="DPI of each image (must be between 70 and 300).")
    parser.add_argument("-l", "--limit_images", type=int, default=0,help="Limit the number of images to process. Use 0 for no limit.")
    parser.add_argument("-s", "--height_scale", type=float, default=0.8,help="Height scale (must be between 0.1 and 1.5).")
    parser.add_argument("-i", "--input", type=str,default=os.getcwd(), help="Input folder path for images.")
    args = parser.parse_args()

    # **Validate DPI**
    if not (70 <= args.dpi <= 300):
        print("ERROR: DPI must be between 70 and 300.")
        return

    # **Validate height scale**
    if not (0.1 <= args.height_scale <= 1.5):
        print("ERROR: Height scale must be between 0.1 and 1.5.")
        return

    current_folder = args.input
    images = get_images_from_folder(current_folder)

    if args.limit_images != 0:
        images = images[: args.limit_images]

    if images:
        create_docx_from_images(images, args.cols_per_row,
                                args.rows_per_page, args.height_scale, args.dpi)
    else:
        print("No image files found in the specified folder.")


if __name__ == "__main__":
    main()
