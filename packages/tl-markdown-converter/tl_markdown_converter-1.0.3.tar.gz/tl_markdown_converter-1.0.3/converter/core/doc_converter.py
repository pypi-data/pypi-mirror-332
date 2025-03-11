"""
DOC转换模块

这个模块负责将Markdown文件或HTML文件转换为DOC格式（Microsoft Word文档）。
使用自定义的HTMLToDocxConverter类（基于Python的html.parser.HTMLParser）解析HTML并将其转换为Word文档，
使用python-docx库创建和操作Word文档对象。
"""

import os
import re
import tempfile
import shutil
import json
import html
from pathlib import Path
from typing import Dict, Any, List, Optional, Union, Tuple
from html.parser import HTMLParser
from urllib.parse import urlparse, unquote

from converter.utils.logger import logger
from converter.utils.config import config

# 尝试导入python-docx库
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx库未安装，DOC转换功能不可用")

# 尝试导入markdown库，用于Markdown到HTML的转换
try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    logger.warning("markdown库未安装，Markdown到HTML的转换功能不可用")


class HTMLToDocxConverter(HTMLParser):
    """
    HTML到Word文档转换器
    
    基于Python的html.parser.HTMLParser，解析HTML并将其转换为Word文档。
    """
    
    def __init__(self):
        super().__init__()
        self.doc = Document()
        self.current_paragraph = None
        self.current_run = None
        self.list_stack = []  # 跟踪嵌套列表
        self.table_stack = []  # 跟踪表格
        self.in_strong = False
        self.in_em = False
        self.in_link = False
        self.link_href = ""
        self.image_paths = []  # 跟踪文档中的图片路径
        
        # 设置默认样式
        self.doc.styles['Normal'].font.name = 'Arial'
        self.doc.styles['Normal'].font.size = Pt(11)
        
        # 初始化当前段落
        self.current_paragraph = self.doc.add_paragraph()
    
    def handle_starttag(self, tag, attrs):
        """处理开始标签"""
        attrs_dict = dict(attrs)
        
        if tag == 'p':
            self.current_paragraph = self.doc.add_paragraph()
            # 处理对齐方式
            if 'align' in attrs_dict:
                if attrs_dict['align'] == 'center':
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif attrs_dict['align'] == 'right':
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif attrs_dict['align'] == 'justify':
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        elif tag == 'h1':
            self.current_paragraph = self.doc.add_heading(level=1)
        elif tag == 'h2':
            self.current_paragraph = self.doc.add_heading(level=2)
        elif tag == 'h3':
            self.current_paragraph = self.doc.add_heading(level=3)
        elif tag == 'h4':
            self.current_paragraph = self.doc.add_heading(level=4)
        elif tag == 'h5':
            self.current_paragraph = self.doc.add_heading(level=5)
        elif tag == 'h6':
            self.current_paragraph = self.doc.add_heading(level=6)
        
        elif tag == 'ul' or tag == 'ol':
            # 添加列表到堆栈
            list_type = 'ul' if tag == 'ul' else 'ol'
            level = len(self.list_stack)
            self.list_stack.append({'type': list_type, 'level': level, 'count': 0})
        
        elif tag == 'li':
            # 创建列表项
            if self.list_stack:
                current_list = self.list_stack[-1]
                current_list['count'] += 1
                
                # 创建新段落
                self.current_paragraph = self.doc.add_paragraph()
                
                # 设置缩进
                indent = current_list['level'] * 0.25
                self.current_paragraph.paragraph_format.left_indent = Inches(indent)
                
                # 添加列表标记
                if current_list['type'] == 'ul':
                    self.current_paragraph.add_run('• ')
                else:
                    self.current_paragraph.add_run(f"{current_list['count']}. ")
        
        elif tag == 'strong' or tag == 'b':
            self.in_strong = True
        
        elif tag == 'em' or tag == 'i':
            self.in_em = True
        
        elif tag == 'a':
            self.in_link = True
            if 'href' in attrs_dict:
                self.link_href = attrs_dict['href']
        
        elif tag == 'br':
            if self.current_paragraph:
                self.current_paragraph.add_run('\n')
        
        elif tag == 'img':
            if 'src' in attrs_dict:
                src = attrs_dict['src']
                alt = attrs_dict.get('alt', '')
                self._add_image(src, alt)
        
        elif tag == 'table':
            # 开始一个新表格
            self.table_stack.append({
                'rows': 0,
                'cols': 0,
                'current_row': -1,
                'current_col': 0,
                'data': []
            })
        
        elif tag == 'tr':
            # 新行
            if self.table_stack:
                table = self.table_stack[-1]
                table['current_row'] += 1
                table['current_col'] = 0
                table['rows'] += 1
                table['data'].append([])
        
        elif tag == 'td' or tag == 'th':
            # 新单元格
            if self.table_stack:
                table = self.table_stack[-1]
                if table['current_row'] >= 0:
                    table['current_col'] += 1
                    if table['current_col'] > table['cols']:
                        table['cols'] = table['current_col']
                    # 添加空单元格数据
                    table['data'][table['current_row']].append('')
                    # 创建临时段落来收集单元格内容
                    self.current_paragraph = self.doc.add_paragraph()
    
    def handle_endtag(self, tag):
        """处理结束标签"""
        if tag == 'p':
            self.current_paragraph = None
            self.current_run = None
        elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.current_paragraph = None
            self.current_run = None
        elif tag == 'strong' or tag == 'b':
            self.in_strong = False
        elif tag == 'em' or tag == 'i':
            self.in_em = False
        elif tag == 'a':
            self.in_link = False
            self.link_href = ""
        elif tag == 'table':
            # 完成表格处理
            if self.table_stack:
                table_info = self.table_stack[-1]
                if table_info['rows'] > 0 and table_info['cols'] > 0:
                    # 创建表格
                    table = self.doc.add_table(rows=table_info['rows'], cols=table_info['cols'])
                    table.style = 'Table Grid'
                    
                    # 填充表格内容
                    for i in range(table_info['rows']):
                        for j in range(table_info['cols']):
                            if i < len(table_info['data']) and j < len(table_info['data'][i]):
                                cell = table.cell(i, j)
                                cell.text = table_info['data'][i][j]
                
                # 移除表格信息
                self.table_stack.pop()
        elif tag == 'tr':
            if self.table_stack:
                self.table_stack[-1]['current_row'] = -1
        elif tag == 'td' or tag == 'th':
            self.current_paragraph = None
            self.current_run = None
        elif tag == 'ul' or tag == 'ol':
            # 结束列表
            if self.list_stack:
                self.list_stack.pop()
        elif tag == 'li':
            # 结束列表项
            self.current_paragraph = None
            self.current_run = None
    
    def handle_data(self, data):
        """处理文本数据"""
        if not data.strip():
            return  # 忽略空白文本
        
        # 确保有当前段落
        if self.current_paragraph is None:
            self.current_paragraph = self.doc.add_paragraph()
        
        # 添加文本
        if self.table_stack and self.table_stack[-1]['current_row'] >= 0:
            # 表格单元格内容
            row = self.table_stack[-1]['current_row']
            col = self.table_stack[-1]['current_col'] - 1
            if row >= 0 and col >= 0:
                self.table_stack[-1]['data'][row][col] += data
        elif self.in_link:
            # 链接文本
            self.current_run = self.current_paragraph.add_run(data)
            self.current_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
            self.current_run.font.underline = True
            # 存储链接信息以便后续处理
            self._add_hyperlink(self.current_paragraph, data, self.link_href)
        elif self.in_strong:
            # 粗体文本
            self.current_run = self.current_paragraph.add_run(data)
            self.current_run.font.bold = True
        elif self.in_em:
            # 斜体文本
            self.current_run = self.current_paragraph.add_run(data)
            self.current_run.font.italic = True
        else:
            # 普通文本
            self.current_run = self.current_paragraph.add_run(data)
    
    def _add_image(self, src, alt):
        """处理图片标签"""
        if not src:
            return
        
        try:
            # 解析图片路径
            img_path = self._resolve_path(src)
            
            # 处理宽度和高度属性
            width = None
            height = None
            
            # 添加图片
            if self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph()
            
            self.current_run = self.current_paragraph.add_run()
            self.current_run.add_picture(img_path, width=width, height=height)
            
            # 记录图片路径
            self.image_paths.append(img_path)
            
            # 添加图片说明
            if alt:
                caption_paragraph = self.doc.add_paragraph(alt)
                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_paragraph.style = 'Caption'
        except Exception as e:
            logger.error(f"添加图片时出错: {str(e)}")
    
    def _resolve_path(self, path):
        """
        解析路径，支持相对路径和绝对路径
        
        Args:
            path (str): 原始路径
            
        Returns:
            str: 解析后的绝对路径
        """
        # 如果是URL，直接返回
        if path.startswith(('http://', 'https://')):
            return path
        
        # 相对路径
        if not os.path.isabs(path):
            base_dir = os.path.dirname(os.path.abspath(__file__))
            return os.path.join(base_dir, path)
        
        # 绝对路径
        return path
        
    def _add_hyperlink(self, paragraph, text, url):
        """添加超链接"""
        # 这个函数使用了python-docx的底层API来添加超链接
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        # 创建超链接XML元素
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # 创建文本运行
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # 添加样式
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        # 添加下划线
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        
        # 添加文本
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        
        return hyperlink


def check_docx_available() -> bool:
    """
    检查python-docx库是否可用
    
    Returns:
        bool: 如果python-docx库可用，返回True，否则返回False
    """
    return DOCX_AVAILABLE


def convert_html_to_doc(
    html_file_path: str, 
    output_dir: Optional[str] = None, 
    options: Optional[Dict[str, Any]] = None
) -> Optional[str]:
    """
    将HTML文件转换为DOC文档
    
    Args:
        html_file_path (str): HTML文件路径
        output_dir (Optional[str], optional): 输出目录. Defaults to None.
        options (Optional[Dict[str, Any]], optional): 转换选项. Defaults to None.
    
    Returns:
        Optional[str]: 输出文件路径，如果转换失败则返回None
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx库不可用，无法进行DOC转换")
        return None

    # 检查文件是否存在
    if not os.path.exists(html_file_path):
        logger.error(f"HTML文件不存在: {html_file_path}")
        return None

    # 解析选项
    options = options or {}
    
    # 检查是否提供了自定义输出文件路径
    custom_output_path = options.get("output_file_path")
    
    try:
        html_path = Path(html_file_path)
        
        # 准备输出目录
        if custom_output_path:
            # 使用自定义输出路径
            output_file_path = Path(custom_output_path)
            # 确保输出目录存在
            output_file_path.parent.mkdir(parents=True, exist_ok=True)
        else:
            # 使用默认输出路径逻辑
            if output_dir:
                output_path = Path(output_dir)
            else:
                output_path = Path(config.get("paths.doc_dir", "output/DOC"))
                
            # 确保输出目录存在
            os.makedirs(output_path, exist_ok=True)
            
            # 获取原始文件名（不含扩展名）
            html_filename = Path(html_file_path).name
            if html_filename.startswith('tmp') and html_path.stem != Path(html_file_path).stem:
                # 这是一个临时文件，尝试从内容中提取原始名称
                try:
                    with open(html_file_path, 'r', encoding='utf-8') as f:
                        html_content = f.read()
                        # 尝试从HTML中获取标题
                        title_match = re.search(r'<title>(.*?)</title>', html_content)
                        if title_match and title_match.group(1):
                            original_name = title_match.group(1)
                            # 清理名称以确保文件名安全
                            original_name = re.sub(r'[\\/*?:"<>|]', "", original_name)
                            # 使用此作为输出文件名
                            output_file_name = original_name + "." + config.get("doc.default_format", "docx")
                        else:
                            # 回退到原始stem名称，不带tmp前缀
                            output_file_name = html_path.stem + "." + config.get("doc.default_format", "docx")
                    
                except Exception as e:
                    logger.warning(f"从HTML中提取原始名称失败: {e}")
                    # 回退到原始名称
                    output_file_name = html_path.stem + "." + config.get("doc.default_format", "docx")
            else:
                # 常规HTML文件，使用其名称
                output_file_name = html_path.stem + "." + config.get("doc.default_format", "docx")
            
            output_file_path = output_path / output_file_name
        
        # 获取默认选项
        default_options = config.get("doc.options", {})
        
        # 合并用户选项
        doc_options = default_options.copy()
        if options:
            doc_options.update(options)
        
        logger.info(f"正在将 {html_file_path} 转换为DOC格式")
        
        # 读取HTML文件内容
        with open(html_file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # 获取HTML文件的基础路径，用于解析相对路径
        base_path = os.path.dirname(os.path.abspath(html_file_path))
        
        # 创建转换器
        converter = HTMLToDocxConverter()
        
        # 解析HTML
        converter.feed(html_content)
        
        # 保存文档
        converter.doc.save(str(output_file_path))
        
        logger.info(f"DOC转换成功，输出文件: {output_file_path}")
        return str(output_file_path)
    
    except Exception as e:
        logger.error(f"转换DOC时出错: {str(e)}")
        import traceback
        logger.debug(traceback.format_exc())
        return None


def convert_md_to_doc(md_file_path, output_dir, doc_options=None):
    """
    将Markdown文件转换为DOC格式

    Args:
        md_file_path (str): Markdown文件路径
        output_dir (str): 输出目录
        doc_options (dict, optional): DOC转换选项. Defaults to None.

    Returns:
        Optional[str]: 输出的DOC文件路径，失败则返回None
    """
    # 检查模块可用性
    if not DOCX_AVAILABLE:
        logger.error("python-docx模块不可用，无法转换为DOC格式")
        return None

    # 检查文件是否存在
    if not os.path.exists(md_file_path):
        logger.error(f"Markdown文件不存在: {md_file_path}")
        return None

    # 获取原始文件名（不含扩展名）
    md_filename = os.path.basename(md_file_path)
    original_name = os.path.splitext(md_filename)[0]

    try:
        # 先转换为临时HTML文件
        with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as temp:
            temp_html_path = temp.name
        
        # 为临时HTML文件添加原始文件信息以便在后续转换中提取
        with open(md_file_path, 'r', encoding='utf-8') as md_file:
            md_content = md_file.read()

        html_content = markdown.markdown(md_content, extensions=config.get("markdown.extensions", ["tables", "fenced_code"]))
        
        # 创建一个包含原始文件名的HTML标题，以便后续转换可以提取
        html_with_title = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>{original_name}</title>
</head>
<body>
{html_content}
</body>
</html>
"""
        # 写入临时HTML文件
        with open(temp_html_path, 'w', encoding='utf-8') as f:
            f.write(html_with_title)
            
        # 直接创建输出文件路径，而不是依赖convert_html_to_doc来生成
        if output_dir:
            output_path = Path(output_dir)
        else:
            output_path = Path(config.get("paths.doc_dir", "output/DOC"))
            
        # 确保输出目录存在
        os.makedirs(output_path, exist_ok=True)
        
        # 使用原始文件名创建输出文件路径
        output_file_name = original_name + "." + config.get("doc.default_format", "docx")
        output_file_path = output_path / output_file_name
        
        # 转换HTML为DOC，但传递预定义的输出路径
        custom_options = doc_options.copy() if doc_options else {}
        custom_options["output_file_path"] = str(output_file_path)
        
        # 转换HTML为DOC
        result = convert_html_to_doc(temp_html_path, output_dir, custom_options)
        
        # 如果转换成功但返回的文件名不是我们期望的，重命名文件
        if result and os.path.exists(result) and result != str(output_file_path) and os.path.exists(str(output_file_path)):
            # 文件已经以正确的名称创建，不需要重命名
            result = str(output_file_path)
        elif result and os.path.exists(result) and result != str(output_file_path):
            # 需要重命名文件
            try:
                os.rename(result, str(output_file_path))
                result = str(output_file_path)
                logger.info(f"重命名DOC文件: {result}")
            except Exception as e:
                logger.warning(f"重命名DOC文件失败: {e}")
        
        # 删除临时HTML文件
        try:
            os.unlink(temp_html_path)
        except:
            pass
            
        return result
    except Exception as e:
        logger.error(f"转换Markdown到DOC失败: {e}")
        return None 