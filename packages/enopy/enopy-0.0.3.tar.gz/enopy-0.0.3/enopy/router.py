import re
import tempfile
from pathlib import Path, PurePosixPath

from fastapi import APIRouter, HTTPException
from starlette.responses import FileResponse
import docx
from fans.logger import get_logger

from eno.lang.call_py import call_py
from eno.lang.call_cl import call_cl
from eno.parse import parse_eno
from eno.env import env


app = APIRouter()
logger = get_logger(__name__)


@app.get('/api/eno/{path:path}')
def api_get_eno(path):
    env.ensure_conf()
    path = follow_alias(PurePosixPath(path), env.aliases)
    target = find_eno_target(path, suffixes = ['.eno', '.js', '.py', '.yaml'])
    if target:
        match target['type']:
            case 'eno':
                return get_file_eno(target['path'])
            case 'py':
                return make_py_eno(target['path'])
            case 'js':
                return make_js_eno(target['path'])
            case 'dir':
                return get_dir_eno(target['path'])
            case _:
                logger.warning(f'invalid target: {target}')
    raise HTTPException(404)


@app.get('/api/note/notes')
def list_notes():
    return list_tree(env.root)


@app.post('/api/eno/py')
def eno_py(req: dict):
    return call_py(req)


@app.post('/api/eno/cl')
def eno_cl(req: dict):
    return call_cl(req)


@app.post('/api/eno/download')
def eno_download(req: dict):
    author = '文川'

    suffix = req['suffix']
    fd, path = tempfile.mkstemp(suffix = f'.{suffix}')
    if req['type'] == 'novel':
        novel = req['data']
        heading = novel['heading']
        sections = novel['sections']
        if suffix == 'txt':
            with open(path, 'w') as f:
                #f.write(f'《{heading}》 作者：{author}\n')
                f.write(f'《{heading}》\n')
                for section in novel['sections']:
                    if len(sections) > 1:
                        f.write(f'\n{section["order"]}. {section["title"] or ""}\n')
                    for para in section['paragraphs']:
                        f.write(f'    {para["text"]}\n')
        elif suffix == 'docx':
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn

            #font_size = Pt(10.5) # 五号
            font_size = Pt(12) # 五号

            doc = docx.Document()
            doc.core_properties.author = author
            doc.core_properties.comments = ''
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            doc.styles['Normal'].font.size = font_size

            para = doc.add_heading('', level = 1)
            run = para.add_run(heading)
            run.font.name = '黑体'
            run.font.size = Pt(15) # 小三号
            run.font.color.rgb = RGBColor(0, 0, 0)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            #para = doc.add_paragraph(f'作者：{author}')
            #para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            for section in sections:
                if len(sections) > 1:
                    para = doc.add_heading('', level = 4)
                    para.paragraph_format.space_after = font_size // 2
                    run = para.add_run(f'{section["order"]}. {section["title"] or ""}')
                    run.font.name = '黑体'
                    run.font.size = font_size
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                for para in section['paragraphs']:
                    p = doc.add_paragraph()
                    run = p.add_run(para['text'])
                    p.style.font.size = font_size
                    if 'style' in para:
                        if para['style'].get('fontWeight') == 'bold':
                            run.bold= True
                            #run.font.name = '黑体'
                            #run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        elif para['style'].get('fontStyle') == 'italic':
                            run.italic= True
                    p.paragraph_format.first_line_indent = font_size * 2
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = font_size // 2

            doc.save(path)
        return FileResponse(
            path,
            filename = req['filename'],
        )



def list_tree(root_path, name = 'root'):
    prefix_len = len(str(env.root))
    dirs = []
    files = []
    for path in root_path.iterdir():
        if ignore_name_regex.match(path.name):
            continue
        if path.is_dir():
            dirs.append(list_tree(path, path.name))
        elif path.suffix in supported_suffixes:
            files.append({
                'name': path.name,
                'path': str(path)[prefix_len:],
                'isfile': True,
            })
    path = str(root_path)[prefix_len:]
    return {
        'name': name,
        'path': path if path else '/',
        'dirs': sorted(dirs, key = lambda d: d['path']),
        'files': sorted(files, key = lambda f: f['path']),
    }


def follow_alias(path: PurePosixPath, aliases: dict[PurePosixPath, PurePosixPath]):
    while path in aliases:
        path = aliases[path]
    return path


def find_eno_target(
        path: PurePosixPath,
        suffixes: list,
        is_path_exists = lambda path: (env.root / path).exists(),
        is_file = lambda path: (env.root / path).is_file(),
        is_dir = lambda path: (env.root / path).is_dir(),
):
    """
    path - eno path relative to eno root
    suffixes - supported suffixes in descending priority order
    """
    if is_path_exists(path):
        if is_file(path):
            return {'type': path.suffix.lstrip('.'), 'path': path}
        elif is_dir(path):
            return {'type': 'dir', 'path': path}
        else:
            return None # not file or dir

    if path.suffix in suffixes:
        path = path.parent / path.stem
    else:
        for suffix in suffixes:
            tpath = path.parent / f'{path.name}{suffix}'
            if is_path_exists(tpath) and is_file(tpath):
                return {'type': suffix.lstrip('.'), 'path': tpath}

    if is_path_exists(path) and is_dir(path):
        return {'type': 'dir', 'path': path}


def get_file_eno(path: PurePosixPath):
    fpath = env.root / path
    with fpath.open() as f:
        text = f.read()
    return make_text_eno(text, path)


def get_dir_eno(path: PurePosixPath):
    dirpath = env.root / path

    def handle(dirpath, suffix, make_ret):
        for name in ['main', dirpath.name]:
            fpath = dirpath / f'{name}{suffix}'
            if fpath.exists():
                return make_ret(fpath.relative_to(env.root))

    return (
        handle(dirpath, '.eno', get_file_eno) or
        handle(dirpath, '.js', make_js_eno) or
        handle(dirpath, '.py', make_py_eno)
    )


def make_text_eno(text, path):
    eno = parse_eno(text, path = '/' / path)
    return {
        'eno': eno,
        'path': str(path),
        'text': text,
    }


def make_js_eno(path):
    return make_text_eno(f'`eno.js ./{path.name}`', path.parent / f'{path.stem}.eno')


def make_py_eno(path):
    return make_text_eno(f'`eno.py ./{path.name}`', path.parent / f'{path.stem}.eno')


supported_suffixes = set(['.md', '.eno'])
ignore_name_regex = re.compile('|'.join(f"({expr})" for expr in [
    r'^\..*',
    r'^__pycache__$',
]))
