import json
from docx import Document

def save_to_word(json_data, output_path):
    """
        将JSON中的文本内容保存为Word文档。
        :param json_data: 解析后的JSON数据
        :param output_path: 输出的Word文件路径
        """
    # 提取文本内容
    text_content = []
    transcripts = json_data.get("transcripts", [])
    for transcript in transcripts:
        content = transcript.get("text", "")
        text_content.append(content)

    # 创建Word文档
    doc = Document()
    doc.add_heading("Transcription Content", level=1)

    # 将每个段落分别添加到文档中
    for content in text_content:
        doc.add_paragraph(content)

    doc.save(output_path)
    print(f"文本已保存到 {output_path}")


def save_to_srt(json_data, output_path):
    """
    将JSON中的字幕内容保存为SRT文件。
    :param json_data: 解析后的JSON数据
    :param output_path: 输出的SRT文件路径
    """
    # 提取字幕内容
    srt_content = []
    transcripts = json_data.get("transcripts", [])
    for i, transcript in enumerate(transcripts, start=1):
        sentences = transcript.get("sentences", [])
        for sentence in sentences:
            start_time = sentence.get("begin_time")
            end_time = sentence.get("end_time")
            text = sentence.get("text")
            sentence_id = sentence["sentence_id"]

            # 转换时间戳为SRT格式
            start_time_srt = f"{start_time // 3600000:02d}:{(start_time // 60000) % 60:02d}:{(start_time // 1000) % 60:02d},{start_time % 1000:03d}"
            end_time_srt = f"{end_time // 3600000:02d}:{(end_time // 60000) % 60:02d}:{(end_time // 1000) % 60:02d},{end_time % 1000:03d}"

            # 创建SRT字幕行
            srt_line = f"{sentence_id}\n{start_time_srt} --> {end_time_srt}\n{text}\n\n"
            srt_content.append(srt_line)

    # 写入SRT文件
    with open(output_path, "w", encoding="utf-8-sig") as srt_file:
        srt_file.writelines(srt_content)
    print(f"字幕已保存到 {output_path}")


def process_json_file(json_file_path, word_output_path, srt_output_path):
    """
    处理JSON文件并调用保存函数。
    :param json_file_path: JSON文件路径
    :param word_output_path: 输出Word文件路径
    :param srt_output_path: 输出SRT文件路径
    """
    # 读取JSON文件
    with open(json_file_path, "r", encoding="utf-8") as file:
        json_data = json.load(file)

    # 保存为Word文档
    save_to_word(json_data, word_output_path)

    # 保存为SRT文件
    save_to_srt(json_data, srt_output_path)


# 示例调用
if __name__ == "__main__":
    json_file_path = "孔令秋血流储备分数FFR在分叉病变中的应用.txt"
    word_output_path = "孔令秋血流储备分数FFR在分叉病变中的应用paraformer.docx"
    srt_output_path = "孔令秋血流储备分数FFR在分叉病变中的应用paraformer.srt"

    process_json_file(json_file_path, word_output_path, srt_output_path)